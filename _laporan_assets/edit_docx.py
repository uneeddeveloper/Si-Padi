# -*- coding: utf-8 -*-
"""Apply dosen revisions to 'Template JUKTISI Simpeda.docx'.

Revisions:
 1. Rename Si-Padi -> SIMPEDA everywhere (body, tables, header, footer).
 2. Masyarakat actor now registers/logs in; role ENUM + profile columns.
 3. Insert 3 required figures (Use Case, ERD, LRS) + dashboard warga,
    each with caption, source, and a purpose-explanation paragraph.
 4. Expand requirement specification (functional + non-functional).
 5. Trim tables to "secukupnya" (drop tanggapan table, renumber).
 6. Add warga-dashboard + laporan results explanation.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = r"C:\laragon\www\Si-Padi"
ASSETS = os.path.join(ROOT, "_laporan_assets")
SRC = os.path.join(ROOT, "Template JUKTISI Simpeda.docx")
DST = os.path.join(ROOT, "Template JUKTISI Simpeda.docx")  # overwrite in place

doc = Document(SRC)
TNR = "Times New Roman"


# ---------------------------------------------------------------- helpers
def style_run(run, size=10, bold=False, italic=False, color=None, name=TNR):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), name)


def set_para_text(p, text, size=10, bold=False, italic=False, align=None):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    style_run(run, size, bold, italic)
    if align is not None:
        p.alignment = align
    return p


def find_para(substr):
    for p in doc.paragraphs:
        if substr in p.text:
            return p
    raise RuntimeError("paragraph not found: " + substr)


def replace_runs(container_paras, old, new):
    n = 0
    for p in container_paras:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                n += 1
    return n


class BlockInserter:
    """Insert a sequence of paragraphs right after a reference paragraph."""
    def __init__(self, ref_para):
        self.anchor = ref_para._p
        self.parent = ref_para._parent

    def _new(self):
        p = OxmlElement('w:p')
        self.anchor.addnext(p)
        self.anchor = p
        return Paragraph(p, self.parent)

    def text(self, text, size=10, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=6, sa=6):
        p = self._new()
        run = p.add_run(text)
        style_run(run, size, bold, italic)
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        return p

    def image(self, path, width_in, sb=6, sa=2):
        p = self._new()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        return p

    def caption(self, text):
        p = self._new()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        style_run(run, 10, bold=True)
        return p

    def source(self, text):
        p = self._new()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        style_run(run, 9, italic=True)
        return p


def set_cell(cell, text, size=9, bold=False, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    style_run(run, size, bold)
    if align is not None:
        para.alignment = align


def rebuild_table(tbl, data_rows):
    """Keep header row (row 0); replace all data rows with data_rows."""
    for row in tbl.rows[1:]:
        row._tr.getparent().remove(row._tr)
    for cells in data_rows:
        r = tbl.add_row()
        for i, val in enumerate(cells):
            set_cell(r.cells[i], val, size=9,
                     align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else None)


# ---------------------------------------------------------------- 1. RENAME
total = 0
total += replace_runs(doc.paragraphs, "Si-Padi", "SIMPEDA")
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            total += replace_runs(c.paragraphs, "Si-Padi", "SIMPEDA")
for sec in doc.sections:
    for hf in (sec.header, sec.footer, sec.first_page_header,
               sec.first_page_footer, sec.even_page_header,
               sec.even_page_footer):
        total += replace_runs(hf.paragraphs, "Si-Padi", "SIMPEDA")
        for t in hf.tables:
            for row in t.rows:
                for c in row.cells:
                    total += replace_runs(c.paragraphs, "Si-Padi", "SIMPEDA")
print("renamed runs:", total)


# ---------------------------------------------------------------- 2. METHOD citation
p = find_para("Keseluruhan tahapan ini diarahkan untuk menghasilkan rancangan")
p.add_run(" Penggunaan metode waterfall pada pengembangan perangkat lunak "
          "sejalan dengan penelitian Pratama et al. (2023) yang menerapkan "
          "model waterfall sebagai kerangka pengembangan sistem yang "
          "sistematis dan terukur.").font.name = TNR
# style the just-added run
style_run(p.runs[-1], 10)


# ------------------------------------------------ 3. REQUIREMENT SPEC (2.2)
p = find_para("Analisis kebutuhan difokuskan pada identifikasi kebutuhan")
set_para_text(
    p,
    "Analisis kebutuhan dilakukan untuk mengidentifikasi secara spesifik "
    "kebutuhan fungsional dan non-fungsional sistem sebelum tahap "
    "perancangan. Kebutuhan fungsional menggambarkan layanan yang wajib "
    "disediakan sistem bagi setiap aktor. Bagi masyarakat, sistem harus "
    "menyediakan fitur registrasi dan login akun, pembuatan pengaduan "
    "disertai unggah foto dan titik lokasi, pelacakan status melalui nomor "
    "tiket, riwayat pengaduan, serta pemberian rating dan ulasan. Bagi "
    "admin, sistem harus mampu mengelola pengaduan, memperbarui status, "
    "memberikan tanggapan, serta mengelola pengumuman dan struktur "
    "organisasi. Bagi superadmin, sistem harus menyediakan manajemen akun "
    "pengguna, pemantauan log aktivitas, dan pengaturan konfigurasi sistem.",
    size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
# add non-functional paragraph after it
bi = BlockInserter(p)
bi.text(
    "Adapun kebutuhan non-fungsional mencakup aspek keamanan (security) "
    "berupa autentikasi, enkripsi kata sandi, serta perlindungan terhadap "
    "CSRF, XSS, dan SQL Injection; aspek skalabilitas (scalability) agar "
    "sistem mampu menangani pertumbuhan jumlah pengaduan; aspek keandalan "
    "(reliability) agar sistem beroperasi secara stabil; serta aspek "
    "kebergunaan (usability) agar antarmuka mudah digunakan oleh masyarakat "
    "awam tanpa pelatihan khusus.")


# ------------------------------------------------ 4. ACTOR (masyarakat login)
p = find_para("Masyarakat memiliki hak akses untuk melakukan input pengaduan")
set_para_text(
    p,
    "Masyarakat memiliki hak akses untuk mendaftarkan akun (registrasi) dan "
    "masuk (login) ke dalam sistem, membuat pengaduan, melacak status "
    "pengaduan melalui nomor tiket, melihat riwayat pengaduan miliknya, "
    "serta memberikan penilaian (rating) dan ulasan terhadap layanan yang "
    "diterima. Admin memiliki akses operasional yang mencakup pengelolaan "
    "pengaduan, pembaruan status laporan, pemberian tanggapan, manajemen "
    "pengumuman, dan manajemen struktur organisasi. Superadmin memiliki hak "
    "akses penuh terhadap seluruh fitur sistem, termasuk manajemen pengguna, "
    "pemantauan log aktivitas, dan pengaturan konfigurasi sistem.",
    size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Figure 1 - Use Case (after actor description)
bi = BlockInserter(p)
bi.image(os.path.join(ASSETS, "usecase.png"), 5.6)
bi.caption("Gambar 1. Use Case Diagram Sistem SIMPEDA")
bi.source("Sumber: Hasil Perancangan Penulis, 2026")
bi.text(
    "Gambar 1 bertujuan untuk menggambarkan interaksi antara ketiga aktor "
    "dengan sistem SIMPEDA. Diagram tersebut menunjukkan bahwa masyarakat "
    "kini harus mendaftar dan login terlebih dahulu sebelum dapat membuat "
    "pengaduan, sehingga setiap laporan terhubung langsung dengan akun "
    "pemiliknya, sedangkan admin dan superadmin menangani sisi pengelolaan "
    "dan administrasi sistem.")


# ------------------------------------------------ 5. SECURITY RBAC (three roles)
p = find_para("membagi pengguna ke dalam dua peran")
set_para_text(
    p,
    "1. Autentikasi berbasis sesi menggunakan guard authentication bawaan "
    "Laravel dengan implementasi RBAC (Role-Based Access Control) yang "
    "membagi pengguna ke dalam tiga peran, yaitu superadmin, admin, dan "
    "masyarakat. Akun masyarakat dibuat melalui fitur registrasi mandiri, "
    "sedangkan akun admin dan superadmin dikelola oleh perangkat desa.",
    size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


# ------------------------------------------------ 6. ERD & LRS figures
p_erd = find_para("ERD berfungsi sebagai representasi grafis")
bi = BlockInserter(p_erd)
bi.image(os.path.join(ASSETS, "erd.png"), 6.3)
bi.caption("Gambar 2. Entity Relationship Diagram (ERD) Sistem SIMPEDA")
bi.source("Sumber: Hasil Perancangan Penulis, 2026")
bi.text(
    "Gambar 2 bertujuan untuk menjelaskan struktur basis data sistem SIMPEDA "
    "beserta relasi antar entitasnya. Entitas users terhubung ke entitas "
    "pengaduans melalui relasi satu-ke-banyak (1:M), sementara setiap "
    "pengaduan dapat memiliki banyak tanggapan namun hanya menerima satu "
    "penilaian (rating), sehingga seluruh data pelayanan pengaduan saling "
    "terintegrasi secara terstruktur.")

p_lrs = find_para("Logical Record Structure (LRS) merupakan representasi")
bi = BlockInserter(p_lrs)
bi.image(os.path.join(ASSETS, "lrs.png"), 6.3)
bi.caption("Gambar 3. Logical Record Structure (LRS) Sistem SIMPEDA")
bi.source("Sumber: Hasil Perancangan Penulis, 2026")
bi.text(
    "Gambar 3 bertujuan untuk memperlihatkan bentuk fisik tabel-tabel basis "
    "data beserta primary key (PK) dan foreign key (FK) yang menghubungkan "
    "antar tabel. LRS menegaskan penerapan integritas referensial melalui "
    "foreign key, misalnya kolom user_id pada tabel pengaduans yang merujuk "
    "ke kolom id pada tabel users.")


# ------------------------------------------------ 7. TABLES (trim + update)
tbls = {}
for t in doc.tables:
    flat = " ".join(c.text for row in t.rows for c in row.cells)
    if "role" in flat and "is_active" in flat:
        tbls["users"] = t
    elif "nomor_tiket" in flat:
        tbls["pengaduans"] = t
    elif "pengirim" in flat:
        tbls["tanggapan"] = t
    elif "Web Server" in flat:
        tbls["spesifikasi"] = t

# users: include nik, role enum w/ masyarakat, nomor_hp, rt_rw
rebuild_table(tbls["users"], [
    ["1", "id", "BIGINT(20) UNSIGNED", "Primary Key, Auto Increment"],
    ["2", "name", "VARCHAR(255)", "Nama lengkap pengguna"],
    ["3", "nik", "VARCHAR(16)", "NIK akun masyarakat (UNIQUE, nullable)"],
    ["4", "email", "VARCHAR(255)", "Alamat email (UNIQUE)"],
    ["5", "password", "VARCHAR(255)", "Kata sandi terenkripsi bcrypt"],
    ["6", "role", "ENUM('superadmin','admin','masyarakat')",
     "Peran pengguna (default 'masyarakat')"],
    ["7", "nomor_hp", "VARCHAR(20)", "Nomor HP warga (nullable)"],
    ["8", "rt_rw", "VARCHAR(10)", "Domisili RT/RW warga (nullable)"],
    ["9", "is_active", "TINYINT(1)", "Status akun (1=aktif, 0=nonaktif)"],
    ["10", "created_at", "TIMESTAMP", "Waktu pembuatan data"],
])

# pengaduans: add user_id FK linking to masyarakat account
rebuild_table(tbls["pengaduans"], [
    ["1", "id", "BIGINT(20) UNSIGNED", "Primary Key, Auto Increment"],
    ["2", "user_id", "BIGINT(20) UNSIGNED",
     "FK -> users(id), pembuat pengaduan (nullable)"],
    ["3", "nomor_tiket", "VARCHAR(20)", "Nomor tiket laporan (UNIQUE)"],
    ["4", "kategori",
     "ENUM('Infrastruktur','Kebersihan','Keamanan','Administrasi',"
     "'Sosial','Lainnya')", "Kategori jenis pengaduan"],
    ["5", "nama_pelapor", "VARCHAR(100)", "Nama masyarakat pelapor"],
    ["6", "urgensi", "ENUM('Rendah','Sedang','Tinggi')",
     "Tingkat urgensi laporan"],
    ["7", "status", "ENUM('Menunggu','Diproses','Selesai','Ditolak')",
     "Status penanganan (default 'Menunggu')"],
    ["8", "judul", "VARCHAR(255)", "Judul pengaduan"],
    ["9", "deskripsi", "TEXT", "Deskripsi lengkap pengaduan"],
    ["10", "foto", "VARCHAR(255)", "Path file foto pendukung (nullable)"],
    ["11", "created_at", "TIMESTAMP", "Waktu pengiriman pengaduan"],
])

# Remove tanggapan table + its caption ("Tabel 3. Struktur Tabel tanggapan...")
if "tanggapan" in tbls:
    cap = find_para("Tabel 3. Struktur Tabel tanggapan_pengaduans")
    cap._p.getparent().remove(cap._p)
    tbls["tanggapan"]._tbl.getparent().remove(tbls["tanggapan"]._tbl)

# Renumber "Tabel 4. Spesifikasi..." -> "Tabel 3."
for p in doc.paragraphs:
    if "Tabel 4. Spesifikasi Perangkat Lunak" in p.text:
        set_para_text(p, "Tabel 3. Spesifikasi Perangkat Lunak",
                      size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        break


# ------------------------------------------------ 8. DASHBOARD WARGA + LAPORAN
p = find_para("Komponen dashboard statistik sistem SIMPEDA meliputi")
bi = BlockInserter(p)
bi.text(
    "Dari sisi masyarakat, sistem SIMPEDA menyediakan dashboard warga yang "
    "menjadi pusat layanan mandiri bagi pelapor. Setelah berhasil login, "
    "warga diarahkan ke halaman beranda yang menampilkan sapaan personal, "
    "tombol aksi cepat (Buat Pengaduan, Lacak Pengaduan, dan Riwayat), serta "
    "daftar riwayat pengaduan miliknya yang dilengkapi nomor tiket, judul "
    "laporan, indikator status berwarna, dan tanggal pengiriman. Tampilan "
    "dashboard warga ditunjukkan pada Gambar 4.")
bi.image(os.path.join(ASSETS, "dashboard_warga.png"), 5.8)
bi.caption("Gambar 4. Tampilan Dashboard Warga Sistem SIMPEDA")
bi.source("Sumber: Implementasi Sistem SIMPEDA, 2026")
bi.text(
    "Gambar 4 bertujuan untuk menunjukkan hasil penelitian pada bagian "
    "antarmuka warga, yaitu bagaimana masyarakat dapat memantau seluruh "
    "pengaduan yang pernah diajukan beserta perkembangan statusnya secara "
    "real-time dalam satu halaman terpadu. Indikator status berwarna "
    "(Menunggu, Diproses, dan Selesai) memudahkan warga memahami posisi "
    "laporan tanpa harus menghubungi kantor desa, sedangkan fitur laporan "
    "rekapitulasi dalam format PDF di sisi admin menjadi keluaran (output) "
    "yang mendukung pengambilan keputusan perangkat desa.")


doc.save(DST)
print("saved:", DST)
